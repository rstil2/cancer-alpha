#!/usr/bin/env python3
"""
Generate Complete Word Document with Embedded Figures and Tables
Creates a publication-ready manuscript with all figures and tables embedded
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import markdown
import re

def create_word_manuscript():
    """Create a complete Word document with embedded figures and tables"""
    
    print("🔄 Generating comprehensive Word manuscript...")
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read the markdown manuscript
    manuscript_path = Path("cancer_alpha_main_manuscript.md")
    if not manuscript_path.exists():
        print(f"❌ Manuscript file not found: {manuscript_path}")
        return None
        
    with open(manuscript_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Add title
    title = "Oncura: A Production-Ready AI System for Multi-Cancer Classification Achieving 95% Balanced Accuracy on Real TCGA Data"
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process each section
    for i, section in enumerate(sections):
        if i == 0:  # Skip title, handle abstract and intro separately
            lines = section.split('\n')
            for line in lines[2:]:  # Skip title line
                if line.startswith('## Abstract'):
                    doc.add_heading('Abstract', 1)
                elif line.startswith('**') and line.endswith('**:'):
                    # Abstract subsection
                    p = doc.add_paragraph()
                    run = p.add_run(line.replace('**', '').replace(':', ':'))
                    run.bold = True
                elif line.strip() and not line.startswith('#'):
                    doc.add_paragraph(line)
        else:
            # Add section heading
            section_lines = section.split('\n')
            section_title = section_lines[0]
            doc.add_heading(section_title, 1)
            
            # Process section content
            in_table = False
            table_lines = []
            
            for line in section_lines[1:]:
                if line.startswith('### '):
                    doc.add_heading(line.replace('### ', ''), 2)
                elif line.startswith('#### '):
                    doc.add_heading(line.replace('#### ', ''), 3)
                elif line.startswith('**Figure') and line.endswith('**'):
                    # Add figure
                    add_figure_to_doc(doc, line)
                elif line.startswith('**Table') and line.endswith('**'):
                    # Add table title
                    p = doc.add_paragraph()
                    run = p.add_run(line.replace('**', ''))
                    run.bold = True
                elif line.startswith('|') and '|' in line:
                    # Table content
                    if not in_table:
                        in_table = True
                        table_lines = []
                    table_lines.append(line)
                elif in_table and not line.startswith('|'):
                    # End of table
                    add_table_to_doc(doc, table_lines)
                    in_table = False
                    table_lines = []
                    if line.strip():
                        doc.add_paragraph(line)
                elif line.strip() and not line.startswith('#'):
                    # Regular paragraph
                    if line.startswith('- ') or line.startswith('* '):
                        # Bullet point
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif re.match(r'^\d+\.', line):
                        # Numbered list
                        doc.add_paragraph(line[line.index('.')+1:].strip(), style='List Number')
                    else:
                        # Regular paragraph
                        p = doc.add_paragraph()
                        add_formatted_text(p, line)
            
            # Handle any remaining table
            if in_table and table_lines:
                add_table_to_doc(doc, table_lines)
    
    # Save document
    output_path = "Cancer_Alpha_Complete_Manuscript.docx"
    doc.save(output_path)
    print(f"✅ Complete manuscript saved as: {output_path}")
    return output_path

def add_figure_to_doc(doc, figure_line):
    """Add figure to document with image if available"""
    
    # Extract figure number and description
    figure_match = re.search(r'Figure (\d+):\s*(.+)', figure_line.replace('**', ''))
    if not figure_match:
        return
    
    fig_num = figure_match.group(1)
    fig_title = figure_match.group(2)
    
    # Add figure title
    p = doc.add_paragraph()
    run = p.add_run(f"Figure {fig_num}: {fig_title}")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Define figure file mappings
    figure_files = {
        '1': 'figure1_model_performance.png',
        '2': 'figure2_cancer_type_performance.png', 
        '3': 'figure3_feature_importance.png',
        '4': 'figure4_comparison_studies.png',
        '5': 'figure5a_confusion_matrix.png',  # Use 5a for Figure 5
        '6': 'figure6_system_architecture.png'
    }
    
    # Add image if file exists
    if fig_num in figure_files:
        fig_path = Path('manuscript_figures') / figure_files[fig_num]
        if fig_path.exists():
            try:
                # Add image
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(fig_path), width=Inches(6))
                
                # Add some space after figure
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️  Could not embed {fig_path}: {e}")
                doc.add_paragraph(f"[Figure {fig_num} - {fig_title}]")
        else:
            print(f"⚠️  Figure file not found: {fig_path}")
            doc.add_paragraph(f"[Figure {fig_num} - {fig_title}]")
    
    # Special handling for Figure 5B (ROC curves)
    if fig_num == '5':
        fig_path_b = Path('manuscript_figures') / 'figure5b_roc_curves.png'
        if fig_path_b.exists():
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(fig_path_b), width=Inches(6))
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️  Could not embed Figure 5B: {e}")

def add_table_to_doc(doc, table_lines):
    """Add formatted table to document"""
    
    if not table_lines:
        return
    
    # Parse table data
    header_line = table_lines[0]
    separator_line = table_lines[1] if len(table_lines) > 1 else ""
    data_lines = table_lines[2:] if len(table_lines) > 2 else []
    
    # Extract headers
    headers = [col.strip() for col in header_line.split('|')[1:-1]]
    
    # Extract data rows
    rows = []
    for line in data_lines:
        if line.strip():
            row = [col.strip() for col in line.split('|')[1:-1]]
            if len(row) == len(headers):  # Only add rows with correct number of columns
                rows.append(row)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = cell_data
    
    # Add space after table
    doc.add_paragraph()

def add_formatted_text(paragraph, text):
    """Add text with basic formatting (bold, italic)"""
    
    # Handle bold text (**text**)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text  
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)

def main():
    """Main function"""
    try:
        # Check if figures directory exists
        if not Path('manuscript_figures').exists():
            print("⚠️  manuscript_figures directory not found!")
            return
        
        # Generate Word document
        output_file = create_word_manuscript()
        
        # Show file info
        file_path = Path(output_file)
        if file_path.exists():
            file_size = file_path.stat().st_size / (1024 * 1024)  # MB
            print(f"📄 File size: {file_size:.1f} MB")
            print(f"📍 Location: {file_path.absolute()}")
            print("\n🎉 Complete manuscript with embedded figures ready for publication!")
        
    except Exception as e:
        print(f"❌ Error generating manuscript: {e}")

if __name__ == "__main__":
    main()
