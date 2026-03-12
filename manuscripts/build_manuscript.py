#!/usr/bin/env python3
"""Generate figures from real pipeline results and build manuscript .docx with embedded figures."""

import csv
import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

RESULTS_DIR = Path(__file__).parent.parent / "data" / "real_model_results"
FIG_DIR = Path(__file__).parent / "figures"
MD_PATH = Path(__file__).parent / "Oncura_Revised_Real_Data.md"
DOCX_PATH = Path(__file__).parent / "Oncura_Revised_Real_Data.docx"
FIG_DIR.mkdir(exist_ok=True)

MODEL_COLORS = {
    "LightGBM": "#2166ac",
    "LogisticRegression": "#4393c3",
    "XGBoost": "#92c5de",
    "RandomForest": "#d1e5f0",
}

MODALITY_COLORS = {
    "Expression": "#1f77b4",
    "Methylation": "#2ca02c",
    "Mutation": "#d62728",
}

CANCER_LABELS = {
    "TCGA-BRCA": "BRCA",
    "TCGA-COAD": "COAD",
    "TCGA-HNSC": "HNSC",
    "TCGA-LIHC": "LIHC",
    "TCGA-LUAD": "LUAD",
    "TCGA-LUSC": "LUSC",
    "TCGA-PRAD": "PRAD",
    "TCGA-STAD": "STAD",
}


def load_results():
    with open(RESULTS_DIR / "model_results.json") as f:
        return json.load(f)


def load_shap():
    features, values = [], []
    with open(RESULTS_DIR / "feature_importance_shap.csv") as f:
        reader = csv.DictReader(f)
        for row in reader:
            features.append(row["feature"])
            values.append(float(row["mean_abs_shap"]))
    return features, values


def model_display_order(results):
    # Keep order with best first for display consistency
    order = ["LightGBM", "LogisticRegression", "XGBoost", "RandomForest"]
    return [m for m in order if m in results]


def fig1_model_comparison(results):
    """Figure 1: Model comparison - CV vs test balanced accuracy."""
    models = model_display_order(results)
    display_names = {
        "LightGBM": "LightGBM",
        "LogisticRegression": "Logistic\nRegression",
        "XGBoost": "XGBoost",
        "RandomForest": "Random\nForest",
    }

    cv_means = [results[m]["cv_mean"] * 100 for m in models]
    cv_stds = [results[m]["cv_std"] * 100 for m in models]
    test_accs = [results[m]["test_balanced_accuracy"] * 100 for m in models]

    x = np.arange(len(models))
    width = 0.35

    fig, ax = plt.subplots(figsize=(8, 5))
    bars1 = ax.bar(
        x - width / 2,
        cv_means,
        width,
        yerr=cv_stds,
        label="CV Balanced Accuracy",
        color="#2166ac",
        edgecolor="white",
        linewidth=0.5,
        capsize=4,
        error_kw={"linewidth": 1.2},
    )
    bars2 = ax.bar(
        x + width / 2,
        test_accs,
        width,
        label="Test Balanced Accuracy",
        color="#b2182b",
        edgecolor="white",
        linewidth=0.5,
    )

    ax.set_ylabel("Balanced Accuracy (%)", fontsize=12, fontweight="bold")
    ax.set_title("Classification Performance Across Models", fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels([display_names[m] for m in models], fontsize=10)
    ax.set_ylim(94, 100.5)
    ax.yaxis.set_major_locator(mticker.MultipleLocator(1))
    ax.legend(fontsize=10, loc="lower left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    for bar in bars1:
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.20,
            f"{bar.get_height():.1f}%",
            ha="center",
            va="bottom",
            fontsize=8.5,
        )
    for bar in bars2:
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.20,
            f"{bar.get_height():.1f}%",
            ha="center",
            va="bottom",
            fontsize=8.5,
        )

    plt.tight_layout()
    plt.savefig(FIG_DIR / "fig1_model_comparison.png", dpi=300, bbox_inches="tight")
    plt.savefig(FIG_DIR / "fig1_model_comparison.tiff", dpi=300, bbox_inches="tight")
    plt.close()
    print("  Figure 1: Model comparison")


def fig2_confusion_matrix(results):
    """Figure 2: LightGBM confusion matrix heatmap."""
    cm = np.array(results["LightGBM"]["confusion_matrix"])
    labels = list(CANCER_LABELS.values())

    fig, ax = plt.subplots(figsize=(7, 6))
    im = ax.imshow(cm, cmap="Blues", aspect="equal")

    ax.set_xticks(range(len(labels)))
    ax.set_yticks(range(len(labels)))
    ax.set_xticklabels(labels, fontsize=10, rotation=45, ha="right")
    ax.set_yticklabels(labels, fontsize=10)
    ax.set_xlabel("Predicted", fontsize=12, fontweight="bold")
    ax.set_ylabel("True", fontsize=12, fontweight="bold")
    ax.set_title("LightGBM Confusion Matrix (Test Set, n=250)", fontsize=13, fontweight="bold")

    for i in range(len(labels)):
        for j in range(len(labels)):
            val = cm[i, j]
            color = "white" if val > cm.max() / 2 else "black"
            ax.text(j, i, str(val), ha="center", va="center", fontsize=11, color=color, fontweight="bold")

    cbar = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.set_label("Sample Count", fontsize=10)

    plt.tight_layout()
    plt.savefig(FIG_DIR / "fig2_confusion_matrix.png", dpi=300, bbox_inches="tight")
    plt.savefig(FIG_DIR / "fig2_confusion_matrix.tiff", dpi=300, bbox_inches="tight")
    plt.close()
    print("  Figure 2: Confusion matrix")


def fig3_per_class_f1(results):
    """Figure 3: Per-class F1-score across all models."""
    models = model_display_order(results)
    classes = list(CANCER_LABELS.keys())
    labels = list(CANCER_LABELS.values())

    fig, ax = plt.subplots(figsize=(10, 5))
    x = np.arange(len(classes))
    width = 0.2

    for i, model in enumerate(models):
        f1_scores = [results[model]["classification_report"][c]["f1-score"] for c in classes]
        offset = (i - 1.5) * width
        ax.bar(
            x + offset,
            f1_scores,
            width,
            label=model.replace("LogisticRegression", "Logistic Regression").replace("RandomForest", "Random Forest"),
            color=MODEL_COLORS[model],
            edgecolor="white",
            linewidth=0.5,
        )

    ax.set_ylabel("F1-Score", fontsize=12, fontweight="bold")
    ax.set_title("Per-Class F1-Score Across Models (Test Set)", fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylim(0.88, 1.02)
    ax.legend(fontsize=9, loc="lower left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.axhline(y=1.0, color="gray", linestyle="--", linewidth=0.5, alpha=0.5)

    plt.tight_layout()
    plt.savefig(FIG_DIR / "fig3_per_class_f1.png", dpi=300, bbox_inches="tight")
    plt.savefig(FIG_DIR / "fig3_per_class_f1.tiff", dpi=300, bbox_inches="tight")
    plt.close()
    print("  Figure 3: Per-class F1")


def feature_modality(feature):
    if feature.startswith("meth_"):
        return "Methylation"
    if feature.startswith("mut_"):
        return "Mutation"
    return "Expression"


def fig4_shap_importance(features, values):
    """Figure 4: Top 20 SHAP feature importance with modality coloring."""
    top_n = 20
    top_features = features[:top_n][::-1]
    top_values = values[:top_n][::-1]
    modalities = [feature_modality(f) for f in top_features]
    colors = [MODALITY_COLORS[m] for m in modalities]

    fig, ax = plt.subplots(figsize=(9, 7))
    ax.barh(range(top_n), top_values, color=colors, edgecolor="white", linewidth=0.5)
    ax.set_yticks(range(top_n))
    ax.set_yticklabels(top_features, fontsize=10)
    ax.set_xlabel("Mean |SHAP Value|", fontsize=12, fontweight="bold")
    ax.set_title("Top 20 Features by SHAP Importance (LightGBM)", fontsize=13, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    from matplotlib.patches import Patch
    legend_elements = [
        Patch(facecolor=MODALITY_COLORS["Expression"], label="Expression"),
        Patch(facecolor=MODALITY_COLORS["Methylation"], label="Methylation"),
        Patch(facecolor=MODALITY_COLORS["Mutation"], label="Mutation"),
    ]
    ax.legend(handles=legend_elements, fontsize=9, loc="lower right", title="Modality", title_fontsize=10)

    plt.tight_layout()
    plt.savefig(FIG_DIR / "fig4_shap_importance.png", dpi=300, bbox_inches="tight")
    plt.savefig(FIG_DIR / "fig4_shap_importance.tiff", dpi=300, bbox_inches="tight")
    plt.close()
    print("  Figure 4: SHAP importance")


def parse_markdown_to_docx(doc):
    """Minimal markdown parser for headings, paragraphs, bullets, and tables."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules and blank lines
        if line.strip() in {"---", ""}:
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue

        # Table
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            headers = [h.strip() for h in line.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = h
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9.5)
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = val
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        # Bullets
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(11)
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(11)
        i += 1


def insert_figures(doc):
    """Insert generated figures under Results section with captions."""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    figure_blocks = [
        ("fig1_model_comparison.png", "Figure 1. Classification performance across models (CV and held-out test balanced accuracy)."),
        ("fig2_confusion_matrix.png", "Figure 2. LightGBM confusion matrix on the held-out test set (n=250)."),
        ("fig3_per_class_f1.png", "Figure 3. Per-class F1-score across all four models."),
        ("fig4_shap_importance.png", "Figure 4. Top 20 SHAP features for LightGBM, colored by data modality."),
    ]

    for fname, caption in figure_blocks:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(FIG_DIR / fname), width=Inches(5.8))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        cap.paragraph_format.space_after = Pt(10)


def build_docx():
    """Build docx from updated markdown and append embedded figures."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    parse_markdown_to_docx(doc)

    # Append figures section near end (keeps it robust without deep AST edits)
    doc.add_page_break()
    h = doc.add_heading("Figures", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    insert_figures(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nManuscript saved: {DOCX_PATH}")


if __name__ == "__main__":
    print("Generating figures...")
    results = load_results()
    features, values = load_shap()

    fig1_model_comparison(results)
    fig2_confusion_matrix(results)
    fig3_per_class_f1(results)
    fig4_shap_importance(features, values)

    print("\nBuilding manuscript .docx...")
    build_docx()
    print("Done.")
