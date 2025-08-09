#!/usr/bin/env python3
"""
Simple Figure Embedding for Word Document
Debug version to properly embed figures
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

def embed_figures():
    """Embed all figures into Word document"""
    
    # Load the Word document
    doc_path = "cancer_alpha_main_manuscript.docx"
    if not os.path.exists(doc_path):
        print(f"❌ Document not found: {doc_path}")
        return
    
    print(f"📄 Loading Word document: {doc_path}")
    doc = Document(doc_path)
    
    # Check figures directory
    fig_dir = "manuscript_figures"
    if not os.path.exists(fig_dir):
        print(f"❌ Figures directory not found: {fig_dir}")
        return
    
    # List all available figures
    figures = {
        'figure1_model_performance.png': 'Model Performance Results',
        'figure2_cancer_type_performance.png': 'Cancer Type-Specific Performance',
        'figure3_feature_importance.png': 'Feature Importance Analysis',
        'figure3b_shap_interpretability.png': 'SHAP Interpretability Analysis',
        'figure4_comparison_studies.png': 'Comparative Study Results',
        'figure5a_confusion_matrix.png': 'Confusion Matrix',
        'figure5b_roc_curves.png': 'ROC Curves',
        'figure6_system_architecture.png': 'System Architecture'
    }
    
    print(f"\n📊 Available figures:")
    for fig_file, description in figures.items():
        fig_path = os.path.join(fig_dir, fig_file)
        exists = "✅" if os.path.exists(fig_path) else "❌"
        print(f"  {exists} {fig_file} - {description}")
    
    # Read all paragraphs first to debug
    print(f"\n🔍 Scanning document paragraphs...")
    figure_paragraphs = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text and ('Figure' in text or 'figure' in text):
            print(f"  Paragraph {i}: {text[:100]}...")
            figure_paragraphs.append((i, paragraph, text))
    
    print(f"\n📝 Found {len(figure_paragraphs)} paragraphs mentioning figures")
    
    # Add figures at the end of document for now (simpler approach)
    print(f"\n📸 Adding figures to document...")
    
    # Add a page break and section header
    doc.add_page_break()
    header_p = doc.add_paragraph()
    header_p.add_run("FIGURES").bold = True
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Empty line
    
    figures_added = 0
    for fig_file, description in figures.items():
        fig_path = os.path.join(fig_dir, fig_file)
        
        if os.path.exists(fig_path):
            try:
                # Add figure title
                title_p = doc.add_paragraph()
                title_run = title_p.add_run(f"Figure {figures_added + 1}: {description}")
                title_run.bold = True
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add the image
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(fig_path, width=Inches(6.0))
                
                # Add space after figure
                doc.add_paragraph()
                
                figures_added += 1
                print(f"  ✅ Added {fig_file}")
                
            except Exception as e:
                print(f"  ❌ Failed to add {fig_file}: {e}")
        else:
            print(f"  ⚠️  File not found: {fig_path}")
    
    # Save the document
    output_path = "Cancer_Alpha_Complete_Manuscript_with_Figures.docx"
    doc.save(output_path)
    
    # Check file size
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path) / (1024 * 1024)  # MB
        print(f"\n✅ Document saved successfully!")
        print(f"📍 Location: {os.path.abspath(output_path)}")
        print(f"📄 File size: {file_size:.1f} MB")
        print(f"📊 Figures embedded: {figures_added}")
    else:
        print(f"\n❌ Failed to save document")

if __name__ == "__main__":
    embed_figures()
