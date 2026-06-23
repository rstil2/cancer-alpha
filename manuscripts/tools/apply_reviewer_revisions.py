"""
Apply reviewer revisions to Oncura_Revised_Real_Data_Revised-THIS_IS_IT.docx
Based on Bioinformatics Advances reviewer comments (BIOADV-2026-164).

Changes applied:
1. Section 1.2: Rewrite bold-bullet headers as prose
2. Methods para 34: Strengthen data leakage / preprocessing-within-folds statement
3. Para 64: Remove markdown artifact, clarify cross-modal importance
4. Insert proper Heading 3 for section 4.5 before para 65
5. Para 65 (CUP): Rewrite to clarify this is simulation, not real CUP validation
6. Para 67: Remove embedded '## 4. Discussion' artifact; reframe deep learning comparison
7. Para 73 (4.3): Add ablation analysis note
8. Para 75 (4.4 Limitations): Rewrite bold bullets as prose; insert 2 new paragraphs
   covering batch effects/tumor purity and variance-based feature selection bias
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IN_PATH  = '/Users/stillwell/projects/cancer-alpha/manuscripts/Oncura_Revised_Real_Data_Revised-THIS_IS_IT.docx'
OUT_PATH = '/Users/stillwell/projects/cancer-alpha/manuscripts/Oncura_Revised_Real_Data_Revised-THIS_IS_IT_R2.docx'

doc = Document(IN_PATH)
paras = doc.paragraphs


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def clear_and_set_text(para, text):
    """Replace all paragraph content with a single plain-text run."""
    p = para._p
    for child in list(p):
        if child.tag not in (qn('w:pPr'), qn('w:sectPr')):
            p.remove(child)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)


def insert_paragraph_after(doc, ref_para, text, style_name='Body Text'):
    """Insert a new paragraph with given text and style immediately after ref_para."""
    tmp = doc.add_paragraph(text, style=style_name)
    tmp._p.getparent().remove(tmp._p)
    ref_para._p.addnext(tmp._p)
    return tmp


# ---------------------------------------------------------------------------
# Capture paragraph references before any insertions
# ---------------------------------------------------------------------------

p12 = paras[12]   # Section 1.2 intro sentence
p13 = paras[13]   # Bold bullet: Single-modality reliance
p14 = paras[14]   # Bold bullet: Class imbalance
p15 = paras[15]   # Bold bullet: Limited biological validation
p34 = paras[34]   # Leakage stress test
p64 = paras[64]   # Cross-modal importance + markdown artifact
p65 = paras[65]   # CUP section content
p67 = paras[67]   # Deep learning comparison + ## 4. Discussion artifact
p73 = paras[73]   # 4.3 Multi-Modal Integration body
p75 = paras[75]   # 4.4 Limitations bold bullets


# ---------------------------------------------------------------------------
# CHANGE 1 – Section 1.2: Convert bold-bullet structure to prose
# ---------------------------------------------------------------------------

# Para 12: keep the brief intro sentence as-is (it's already plain prose)
# Para 13: remove "**Single-modality reliance.**" bold header
clear_and_set_text(
    p13,
    "Most classification studies use gene expression alone (5, 6), potentially missing "
    "complementary information encoded in DNA methylation and somatic mutations. DNA "
    "methylation patterns are highly tissue-specific and reflect both developmental "
    "lineage and epigenetic reprogramming in cancer (8, 9). Somatic mutations capture "
    "distinct biology including DNA repair deficiency and oncogene activation (10, 11). "
    "Multi-modal integration combining these data types could capture transcriptional "
    "programs, epigenetic states, and mutational signatures that jointly define cancer "
    "types (12, 13)."
)

# Para 14: remove "**Class imbalance.**" bold header
clear_and_set_text(
    p14,
    "Natural cancer incidence distributions in TCGA are heavily skewed—breast cancer has "
    "over 1,000 samples while rarer types have fewer than 50. The standard remedy is "
    "synthetic oversampling via SMOTE or variants (14), which generates artificial samples "
    "through feature-space interpolation. While effective for numerical balance, synthetic "
    "samples do not represent real patients and may introduce biologically implausible "
    "feature combinations, particularly when interpolating between molecularly distinct "
    "tumor subtypes."
)

# Para 15: remove "**Limited biological validation of model features.**" bold header
clear_and_set_text(
    p15,
    "A third limitation concerns biological validation of model features. Many studies "
    "report classification accuracy without examining whether the features driving "
    "predictions correspond to established cancer biology (15). A model might achieve "
    "high accuracy by learning batch effects or tissue-processing artifacts rather than "
    "genuine oncogenic programs. Systematic validation of feature importance against "
    "known cancer biomarkers is necessary to confirm that a classifier has learned "
    "biologically meaningful patterns."
)

print("CHANGE 1 done: Section 1.2 rewritten as prose")


# ---------------------------------------------------------------------------
# CHANGE 2 – Methods para 34: Strengthen data-leakage / preprocessing statement
# ---------------------------------------------------------------------------

clear_and_set_text(
    p34,
    "To prevent data leakage, all preprocessing steps were applied strictly within each "
    "cross-validation training fold and then applied—without refitting—to the "
    "corresponding validation fold and to the final held-out test set. Specifically: "
    "(1) feature standardization (zero mean, unit variance) was computed on training-fold "
    "samples only and applied to validation and test samples; (2) high-variance feature "
    "selection was performed exclusively on training-fold data during each cross-validation "
    "iteration, never on the complete dataset; and (3) the held-out test set (n=250) "
    "contained no patients present in any form in the training set, verified by "
    "cross-referencing TCGA case submitter IDs across all three modalities prior to "
    "splitting."
)

print("CHANGE 2 done: Data-leakage statement strengthened")


# ---------------------------------------------------------------------------
# CHANGE 3 – Para 64: Remove markdown artifact; refine cross-modal text
# ---------------------------------------------------------------------------

clear_and_set_text(
    p64,
    "For some cancer types, DNA methylation was the primary driver of classification, "
    "whereas for others, gene expression carried most of the SHAP weight, illustrating "
    "the complementary nature of multi-modal integration and providing empirical "
    "justification for including all three data modalities."
)

# Insert a proper Heading 3 for section 4.5 between p64 and p65
insert_paragraph_after(doc, p64, "4.5 Cancer of Unknown Primary: Proof of Concept",
                        style_name='Heading 3')

print("CHANGE 3 done: Para 64 fixed; Heading 3 for section 4.5 inserted")


# ---------------------------------------------------------------------------
# CHANGE 4 – Para 65 (CUP section): Clarify simulation nature, remove
#             overstated clinical-utility language
# ---------------------------------------------------------------------------

clear_and_set_text(
    p65,
    "As a proof-of-concept, we assessed the classifier's ability to assign probable "
    "tissue of origin to held-out TCGA primary tumor samples in a CUP-simulation "
    "scenario. It is important to note that this evaluation used primary tumor samples "
    "with known histological origin rather than genuine clinical CUP specimens. Metastatic "
    "and poorly differentiated tumors may exhibit substantially altered molecular profiles "
    "compared with primary tumors, and TCGA samples were collected at initial diagnosis "
    "rather than from the metastatic setting. Validation on authentic CUP biopsies—as "
    "well as on formalin-fixed paraffin-embedded (FFPE) specimens typical of clinical "
    "workflows—will be required before any claim of clinical utility for CUP diagnosis "
    "can be substantiated (23)."
)

print("CHANGE 4 done: CUP section rewritten to clarify simulation nature")


# ---------------------------------------------------------------------------
# CHANGE 5 – Para 67: Remove '## 4. Discussion' artifact; reframe deep-learning
#             comparison as specific to small-n setting
# ---------------------------------------------------------------------------

clear_and_set_text(
    p67,
    "We selected tree-based ensemble models (LightGBM, XGBoost, Random Forest) for "
    "their interpretability and training stability with structured multi-omic features "
    "in the context of this dataset (n=1,248 integrated samples). It should be noted "
    "that this finding reflects the practical constraints of small-sample, high-"
    "dimensional genomic settings: deep neural networks and transformer-based "
    "architectures are known to require substantially larger training cohorts than "
    "gradient-boosting approaches, and their relative performance may differ considerably "
    "at population scale. Conclusions about general model superiority across genomic "
    "classification tasks should therefore not be drawn from this comparison."
)

print("CHANGE 5 done: Para 67 cleaned and deep-learning framing moderated")


# ---------------------------------------------------------------------------
# CHANGE 6 – Para 73 (Discussion 4.3): Add ablation-analysis note
# ---------------------------------------------------------------------------

current_p73 = p73.text.rstrip()
clear_and_set_text(
    p73,
    current_p73 + " The dominance of expression features in the top SHAP rankings "
    "raises the question of how much predictive value methylation and mutation modalities "
    "contribute independently; a formal ablation study comparing expression-only, "
    "methylation-only, mutation-only, and fully integrated models would quantify these "
    "modality-specific contributions and is a priority for future work."
)

print("CHANGE 6 done: Ablation analysis note added to Discussion 4.3")


# ---------------------------------------------------------------------------
# CHANGE 7 – Para 75 (Limitations 4.4): Rewrite as prose; insert batch-effects
#             and feature-selection-bias paragraphs
# ---------------------------------------------------------------------------

# Replace the bold-bullet limitations paragraph with plain prose
clear_and_set_text(
    p75,
    "The current analysis classifies eight major cancer types—a well-defined task with "
    "substantial TCGA representation—and extension to rare cancers or cancers of unknown "
    "primary would present additional challenges given their lower prevalence in existing "
    "datasets. All training and test data originate from TCGA, and the held-out test set "
    "was drawn from the same source distribution; generalization to independent cohorts "
    "such as ICGC or CPTAC has not been demonstrated. Prospective validation on samples "
    "collected at different institutions, on different sequencing platforms, or from "
    "metastatic and poorly differentiated tumors would be required before any clinical "
    "application is considered."
)

# Insert batch effects / tumor purity paragraph after p75
p75b = insert_paragraph_after(
    doc, p75,
    "Batch effects, tumor purity, and platform-specific variability represent potential "
    "confounders that were not explicitly modeled. TCGA data were collected across "
    "multiple participating institutions with heterogeneous sample preparation and "
    "sequencing protocols; while TCGA harmonization pipelines reduce some sources of "
    "variation, residual center-specific effects may influence classification, "
    "particularly for methylation arrays with known platform dependencies. Tumor purity—"
    "the fraction of neoplastic versus stromal and immune cells in a sample—varies across "
    "samples and cancer types and can affect both expression and methylation measurements "
    "in ways that may differ systematically from clinical biopsy specimens.",
    style_name='Body Text'
)

# Insert variance-based feature selection bias paragraph after p75b
insert_paragraph_after(
    doc, p75b,
    "Variance-based feature selection, while computationally efficient, may "
    "preferentially retain tissue-specific expression markers that discriminate well "
    "among the eight primary tumor types studied here but may generalize less effectively "
    "to metastatic or poorly differentiated tumors. The top-ranked CpG probes additionally "
    "require further annotation to determine their genomic context and regulatory "
    "significance beyond their discriminative value in the current analysis.",
    style_name='Body Text'
)

print("CHANGE 7 done: Limitations section rewritten as prose with expanded content")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT_PATH)
print(f"\nRevised manuscript saved to:\n  {OUT_PATH}")
