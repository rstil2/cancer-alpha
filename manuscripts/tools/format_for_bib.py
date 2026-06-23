"""
Reformat Oncura_Revised_Real_Data_Revised-THIS_IS_IT_R2.docx
for submission to Briefings in Bioinformatics (Oxford, OUP).

Changes applied:
1. In-text citations: (N) / (N, M) / (N-M) → [N] / [N, M] / [N-M]
2. Add Keywords paragraph after abstract
3. Reference list: remove trailing periods from journal abbreviations
4. Remove internal metadata lines at the end (word count, tables, etc.)
5. Output: Oncura_BriefingsBioinformatics.docx
"""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IN_PATH  = '/Users/stillwell/projects/cancer-alpha/manuscripts/Oncura_Revised_Real_Data_Revised-THIS_IS_IT_R2.docx'
OUT_PATH = '/Users/stillwell/projects/cancer-alpha/manuscripts/Oncura_BriefingsBioinformatics.docx'

KEYWORDS = ('Keywords: multi-modal cancer classification; TCGA; '
            'machine learning; DNA methylation; SHAP interpretability; '
            'multi-omics integration')

doc = Document(IN_PATH)
paras = doc.paragraphs


# ---------------------------------------------------------------------------
# Helper: replace all text in a paragraph, preserving style but not run fmt
# ---------------------------------------------------------------------------

def replace_para_text(para, new_text):
    p = para._p
    for child in list(p):
        if child.tag not in (qn('w:pPr'), qn('w:sectPr')):
            p.remove(child)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = new_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)


def insert_paragraph_after(doc, ref_para, text, style_name='Body Text'):
    tmp = doc.add_paragraph(text, style=style_name)
    tmp._p.getparent().remove(tmp._p)
    ref_para._p.addnext(tmp._p)
    return tmp


# ---------------------------------------------------------------------------
# 1. Convert in-text citations from (N, M) to [N, M] in every paragraph
#    Handles: (5), (5, 6), (5-8), (5, 6, 7), etc.
# ---------------------------------------------------------------------------

# Regex: parenthesised groups that contain only digits, commas, hyphens, spaces
# e.g. "(5)", "(3, 4)", "(8, 9)", "(12, 13)", "(1–4)"
# Negative lookbehind (?<!\d) prevents matching volume/issue numbers like 28(9)
cite_pattern = re.compile(r'(?<!\d)\((\d[\d,\s\-–]+\d|\d)\)')

changed_count = 0
for para in doc.paragraphs:
    full_text = para.text
    if not full_text.strip():
        continue
    new_text = cite_pattern.sub(lambda m: '[' + m.group(1) + ']', full_text)
    if new_text != full_text:
        replace_para_text(para, new_text)
        changed_count += 1

print(f'Citation format changed in {changed_count} paragraphs')


# ---------------------------------------------------------------------------
# 2. Add Keywords paragraph immediately after the abstract paragraph
#    The abstract is paragraph index 5 (style: First Paragraph, after
#    the "Abstract" heading at index 4).
# ---------------------------------------------------------------------------

# Find the abstract body paragraph by walking through paragraphs sequentially,
# tracking _p elements to avoid stale reference issues.
abstract_p_elem = None
all_p_elems = [p._p for p in doc.paragraphs]   # snapshot of _p elements
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 2' and para.text.strip() == 'Abstract':
        # Look ahead for the first non-empty paragraph after this heading
        for para2 in doc.paragraphs[i + 1:]:
            if para2.text.strip():
                abstract_p_elem = para2._p
                break
        break

if abstract_p_elem is not None:
    matched = next((p for p in doc.paragraphs if p._p is abstract_p_elem), None)
    if matched:
        insert_paragraph_after(doc, matched, KEYWORDS, style_name='Body Text')
        print('Keywords paragraph inserted after abstract')
    else:
        print('WARNING: could not re-find abstract paragraph; keywords not inserted')
else:
    print('WARNING: abstract paragraph not found; keywords not inserted')


# ---------------------------------------------------------------------------
# 3. Fix reference list: remove trailing period from journal abbreviation
#    Pattern: "Nat Med. 2022;" → "Nat Med 2022;"
#    Specifically: <abbrev>. <year>;  where abbrev is not a sentence-ending word
# ---------------------------------------------------------------------------

# Match abbreviated journal names followed by a period before the year
# e.g.  "Nat Med. 2022;"  "BMC Genomics. 2017;"  "Cell. 2018;"  "Science. 2013;"
ref_fix_pattern = re.compile(
    r'([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)*)\.'   # journal abbreviation(s)
    r'(?=\s+\d{4}[;:])'                              # followed by year + separator
)

ref_heading_seen = False
ref_changed = 0
for para in doc.paragraphs:
    txt = para.text.strip()
    if para.style.name == 'Heading 2' and txt == 'References':
        ref_heading_seen = True
        continue
    if ref_heading_seen and txt:
        new_txt = ref_fix_pattern.sub(r'\1', para.text)
        if new_txt != para.text:
            replace_para_text(para, new_txt)
            ref_changed += 1

print(f'Reference abbreviation periods removed in {ref_changed} references')


# ---------------------------------------------------------------------------
# 4. Remove internal metadata lines at the bottom of the manuscript
#    (*Manuscript word count: ~5,800*, *Tables: 3*, etc.)
# ---------------------------------------------------------------------------

metadata_patterns = [
    re.compile(r'^\*?Manuscript word count', re.IGNORECASE),
    re.compile(r'^\*?Tables:', re.IGNORECASE),
    re.compile(r'^\*?Figures:', re.IGNORECASE),
    re.compile(r'^\*?References:', re.IGNORECASE),
]

removed = 0
for para in doc.paragraphs:
    txt = para.text.strip()
    if any(p.match(txt) for p in metadata_patterns):
        # Clear the paragraph text (can't delete paragraphs easily; blank it)
        replace_para_text(para, '')
        removed += 1

print(f'{removed} metadata lines cleared')


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT_PATH)
print(f'\nFormatted manuscript saved to:\n  {OUT_PATH}')
