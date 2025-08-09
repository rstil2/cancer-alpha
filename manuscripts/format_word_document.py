#!/usr/bin/env python3
"""
Format Word Document with Modern Styling
Improves fonts, spacing, and overall appearance
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def format_word_document():
    """Apply modern formatting to the Word document"""
    
    doc_path = "Cancer_Alpha_Complete_Manuscript_with_Figures.docx"
    if not os.path.exists(doc_path):
        print(f"❌ Document not found: {doc_path}")
        return
    
    print(f"🎨 Loading document for formatting: {doc_path}")
    doc = Document(doc_path)
    
    # Define modern fonts in order of preference
    modern_fonts = ['Calibri', 'Arial', 'Helvetica', 'Segoe UI']
    
    print(f"📝 Applying modern formatting...")
    
    # Update document styles
    update_document_styles(doc, modern_fonts)
    
    # Format all paragraphs
    format_paragraphs(doc, modern_fonts)
    
    # Format tables if any
    format_tables(doc, modern_fonts)
    
    # Save the formatted document
    output_path = "Cancer_Alpha_Complete_Manuscript_FORMATTED.docx"
    doc.save(output_path)
    
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path) / (1024 * 1024)  # MB
        print(f"\n✅ Formatted document saved!")
        print(f"📍 Location: {os.path.abspath(output_path)}")
        print(f"📄 File size: {file_size:.1f} MB")
        print(f"🎨 Applied modern Calibri/Arial formatting")
        print(f"📏 Improved spacing and typography")
    else:
        print(f"\n❌ Failed to save formatted document")

def update_document_styles(doc, fonts):
    """Update document-wide styles"""
    
    # Get or create Normal style
    try:
        normal_style = doc.styles['Normal']
    except KeyError:
        normal_style = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    
    # Set font
    font = normal_style.font
    font.name = fonts[0]  # Calibri as primary choice
    font.size = Pt(11)
    
    # Paragraph formatting
    paragraph_format = normal_style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Update heading styles
    for i in range(1, 4):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = fonts[0]
            heading_font.bold = True
            
            if i == 1:
                heading_font.size = Pt(16)
                heading_font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # Professional blue
            elif i == 2:
                heading_font.size = Pt(14)
                heading_font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)
            else:
                heading_font.size = Pt(12)
                heading_font.color.rgb = RGBColor(0x4F, 0x6F, 0x9F)
                
            # Heading spacing
            heading_para = heading_style.paragraph_format
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)
            
        except KeyError:
            pass  # Style doesn't exist, skip

def format_paragraphs(doc, fonts):
    """Format all paragraphs with modern styling"""
    
    formatted_count = 0
    
    for paragraph in doc.paragraphs:
        # Skip empty paragraphs
        if not paragraph.text.strip():
            continue
            
        # Format each run in the paragraph
        for run in paragraph.runs:
            if run.font.name != fonts[0]:
                run.font.name = fonts[0]
                
            # Ensure readable font size
            if run.font.size is None or run.font.size < Pt(10):
                run.font.size = Pt(11)
        
        # Paragraph formatting
        para_format = paragraph.paragraph_format
        
        # Detect heading-like paragraphs (bold, larger, or starting with specific patterns)
        is_heading = False
        if paragraph.runs and paragraph.runs[0].bold:
            text = paragraph.text.strip()
            if (text.startswith(('Abstract', 'Introduction', 'Methods', 'Results', 'Discussion', 
                               'Conclusion', 'References', 'FIGURES', 'Table', 'Figure')) or
                text.isupper() or
                any(run.font.size and run.font.size >= Pt(14) for run in paragraph.runs)):
                is_heading = True
        
        if is_heading:
            # Heading formatting
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(12)
            para_format.keep_with_next = True
            
            # Make sure it's bold and properly sized
            for run in paragraph.runs:
                run.bold = True
                if run.font.size is None or run.font.size < Pt(12):
                    run.font.size = Pt(14)
        else:
            # Body text formatting
            para_format.space_after = Pt(8)
            para_format.line_spacing = 1.15
            
            # Justify text for better appearance
            if len(paragraph.text.strip()) > 50:  # Only justify longer paragraphs
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        formatted_count += 1
    
    print(f"  ✅ Formatted {formatted_count} paragraphs")

def format_tables(doc, fonts):
    """Format tables with modern styling"""
    
    table_count = 0
    
    for table in doc.tables:
        table_count += 1
        
        # Table-wide formatting
        table.style = 'Light Grid Accent 1'
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = fonts[0]
                        run.font.size = Pt(10)
                    
                    # Cell spacing
                    paragraph.paragraph_format.space_after = Pt(3)
    
    if table_count > 0:
        print(f"  ✅ Formatted {table_count} tables")

def set_document_margins(doc):
    """Set professional document margins"""
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

if __name__ == "__main__":
    format_word_document()
