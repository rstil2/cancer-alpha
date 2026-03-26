#!/usr/bin/env python3
"""Build Combined_Manuscript_SciReports_2026.docx with all 6 figures embedded."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

BASE = Path(__file__).parent
SCIENCE = BASE.parent / "science"
EXPERIMENTS = BASE.parent / "experiments"
MD_PATH = BASE / "Combined_Manuscript_SciReports_2026.md"
DOCX_PATH = BASE / "Combined_Manuscript_SciReports_2026.docx"

# Map each figure to its component images and captions
# Combined manuscript has 6 figures, some with A/B panels
FIGURES = [
    {
        "number": 1,
        "caption": (
            "Figure 1. Study design and data flow. Multi-modal features "
            "(gene expression, DNA methylation, somatic mutations) from TCGA "
            "are integrated and evaluated across two sample regimes "
            "(full-cohort n=1,248 and minimal-data n=158), with CUP validation "
            "and cancer subtype prediction."
        ),
        "panels": [
            SCIENCE / "figures" / "figure1_conceptual.png",
        ],
    },
    {
        "number": 2,
        "caption": (
            "Figure 2. Full-cohort classification performance. "
            "(A) Balanced accuracy for four classifiers on cross-validation "
            "and held-out test sets. (B) Confusion matrix for LightGBM on "
            "the held-out test set (n=250)."
        ),
        "panels": [
            BASE / "figures" / "fig1_model_comparison.png",
            BASE / "figures" / "fig2_confusion_matrix.png",
        ],
    },
    {
        "number": 3,
        "caption": (
            "Figure 3. Model complexity versus performance in the minimal-data "
            "setting (n=158). (A) Balanced accuracy versus number of trainable "
            "parameters across 12 architectures, showing inverse correlation "
            "(R\u00b2=0.78). (B) Learning curves demonstrating LightGBM "
            "convergence at n=75 versus transformers requiring n>200."
        ),
        "panels": [
            SCIENCE / "figures" / "figure1_complexity.png",
            SCIENCE / "figures" / "figure2_learning_curves.png",
        ],
    },
    {
        "number": 4,
        "caption": (
            "Figure 4. CUP validation results. (A) Per-cancer-type accuracy "
            "across 10 repeated evaluations (2,500 total predictions). "
            "(B) Calibration analysis: accuracy versus confidence for "
            "high-confidence predictions."
        ),
        "panels": [
            EXPERIMENTS / "cup_results" / "cup_validation_accuracy.png",
            EXPERIMENTS / "cup_results" / "cup_calibration_analysis.png",
        ],
    },
    {
        "number": 5,
        "caption": (
            "Figure 5. Cancer subtype prediction. (A) Cross-validated balanced "
            "accuracy for BRCA (5 subtypes), LUAD (3 subtypes), and COAD "
            "(4 subtypes). (B) Top discriminating features per cancer type."
        ),
        "panels": [
            EXPERIMENTS / "subtype_results" / "subtype_cv_accuracy.png",
            EXPERIMENTS / "subtype_results" / "subtype_shap_features.png",
        ],
    },
    {
        "number": 6,
        "caption": (
            "Figure 6. SHAP feature importance. (A) Top 20 features by mean "
            "absolute SHAP value for the full-cohort LightGBM, colored by data "
            "modality. (B) Cancer-type-specific feature patterns showing "
            "tissue-appropriate biomarker identification."
        ),
        "panels": [
            BASE / "figures" / "fig4_shap_importance.png",
        ],
    },
]


def add_inline_formatted_text(paragraph, text):
    """Parse inline markdown bold/italic and add as runs."""
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


def parse_markdown_to_docx(doc):
    """Parse the combined manuscript markdown into a Word document."""
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip blank lines and horizontal rules
        if line.strip() in {"---", ""}:
            i += 1
            continue

        # Skip markdown image references (![...](...))
        if line.strip().startswith("!["):
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue

        # Table
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Separate header, separator, and data rows
            headers = [h.strip() for h in table_lines[0].strip().strip("|").split("|")]
            data_rows = []
            for tl in table_lines[1:]:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                # Skip separator rows (|---|---|)
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                data_rows.append(cells)

            num_cols = len(headers)
            table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, h in enumerate(headers):
                if c < num_cols:
                    cell = table.rows[0].cells[c]
                    cell.text = h
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(9.5)
            for r_idx, row in enumerate(data_rows):
                for c_idx, val in enumerate(row):
                    if c_idx < num_cols:
                        cell = table.rows[r_idx + 1].cells[c_idx]
                        cell.text = val
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        # Numbered list items
        num_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_formatted_text(p, num_match.group(2))
            for r in p.runs:
                r.font.size = Pt(11)
            i += 1
            continue

        # Bullets
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatted_text(p, line[2:])
            for r in p.runs:
                r.font.size = Pt(11)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Italic metadata lines (e.g. *Manuscript word count: ~6,200*)
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Default paragraph — collect continuation lines
        para_text = line.strip()
        while (i + 1 < len(lines)
               and lines[i + 1].strip()
               and not lines[i + 1].startswith("#")
               and not lines[i + 1].startswith("|")
               and not lines[i + 1].startswith("- ")
               and not lines[i + 1].startswith("* ")
               and not lines[i + 1].startswith("> ")
               and not lines[i + 1].strip().startswith("![")
               and not re.match(r"^---+\s*$", lines[i + 1])
               and not re.match(r"^\d+\.\s+", lines[i + 1])):
            i += 1
            para_text += " " + lines[i].strip()

        p = doc.add_paragraph()
        add_inline_formatted_text(p, para_text)
        for r in p.runs:
            r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
        i += 1


def insert_figures(doc):
    """Insert all 6 figures with their panels and captions."""
    doc.add_page_break()
    h = doc.add_heading("Figures", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    for fig in FIGURES:
        # Check all panels exist
        missing = [p for p in fig["panels"] if not p.exists()]
        if missing:
            for m in missing:
                print(f"  WARNING: Missing panel for Figure {fig['number']}: {m}")

        # Insert each panel image
        for panel_path in fig["panels"]:
            if panel_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # Use 5.8in for single panels, 5.0in for multi-panel
                width = Inches(5.8) if len(fig["panels"]) == 1 else Inches(5.0)
                run.add_picture(str(panel_path), width=width)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[Image not found: {panel_path.name}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(fig["caption"])
        run.italic = True
        run.font.size = Pt(10)
        cap.paragraph_format.space_after = Pt(14)

        print(f"  Figure {fig['number']}: {len([p for p in fig['panels'] if p.exists()])}/{len(fig['panels'])} panels embedded")


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Set heading styles
    for level in range(1, 5):
        try:
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = "Times New Roman"
            hs.font.bold = True
            sizes = {1: 14, 2: 13, 3: 12, 4: 11}
            hs.font.size = Pt(sizes.get(level, 11))
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    print("Parsing manuscript text...")
    parse_markdown_to_docx(doc)

    print("Embedding figures...")
    insert_figures(doc)

    doc.save(str(DOCX_PATH))
    size_kb = DOCX_PATH.stat().st_size / 1024
    print(f"\nSaved: {DOCX_PATH} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    build()
