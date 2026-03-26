#!/usr/bin/env python3
"""Convert the three revision markdown files to Word documents."""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn

BASE = "/Users/stillwell/projects/cancer-alpha/manuscripts"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style=None, bold=False, italic=False,
                         font_size=11, alignment=None, space_after=6):
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    return p


def parse_markdown_to_word(md_path, docx_path):
    with open(md_path, "r") as f:
        lines = f.readlines()

    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level in range(1, 5):
        try:
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = "Calibri"
            hs.font.bold = True
            sizes = {1: 16, 2: 14, 3: 12, 4: 11}
            hs.font.size = Pt(sizes.get(level, 11))
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    i = 0
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip pure horizontal rules
        if re.match(r"^---+\s*$", line):
            i += 1
            continue

        # Table handling
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator rows (|---|---|)
            if all(re.match(r"^[-:]+$", c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            # Check if next line continues the table
            if i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                continue
            else:
                # End of table, render it
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                    tbl.style = "Light Grid Accent 1"
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < num_cols:
                                cell = tbl.cell(ri, ci)
                                cell.text = cell_text
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = "Calibri"
                                        run.font.size = Pt(9)
                                    paragraph.paragraph_format.space_after = Pt(2)
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                    doc.add_paragraph()  # spacer
                table_rows = []
                continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown bold from headings
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                run.font.name = "Calibri"
            i += 1
            continue

        # Bold line (like **Keywords:**)
        bold_match = re.match(r"^\*\*(.+?)\*\*\s*(.*)$", line)
        if bold_match and not line.startswith("- ") and not line.startswith("* "):
            p = doc.add_paragraph()
            run = p.add_run(bold_match.group(1))
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            rest = bold_match.group(2)
            if rest:
                run2 = p.add_run(" " + rest)
                run2.font.name = "Calibri"
                run2.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Numbered list items
        num_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if num_match:
            text = num_match.group(2)
            text = process_inline_formatting(doc, text, is_list=True, list_num=num_match.group(1))
            i += 1
            continue

        # Bullet items
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatted_text(p, text)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_text = line.strip()
        while (i + 1 < len(lines)
               and lines[i + 1].strip()
               and not lines[i + 1].startswith("#")
               and not lines[i + 1].startswith("|")
               and not lines[i + 1].startswith("- ")
               and not lines[i + 1].startswith("* ")
               and not lines[i + 1].startswith("> ")
               and not re.match(r"^---+\s*$", lines[i + 1])
               and not re.match(r"^\d+\.\s+", lines[i + 1])
               and not re.match(r"^\*\*(.+?)\*\*\s*$", lines[i + 1])):
            i += 1
            para_text += " " + lines[i].strip()

        p = doc.add_paragraph()
        add_inline_formatted_text(p, para_text)
        for run in p.runs:
            run.font.name = "Calibri"
            if run.font.size is None:
                run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
        i += 1

    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    return size_kb


def add_inline_formatted_text(paragraph, text):
    """Parse inline markdown (bold, italic) and add as runs."""
    # Split on bold and italic markers
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        elif part:
            run = paragraph.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def process_inline_formatting(doc, text, is_list=False, list_num=None):
    """Add a paragraph with inline formatting for numbered lists."""
    if is_list and list_num:
        p = doc.add_paragraph(style="List Number")
    else:
        p = doc.add_paragraph()
    add_inline_formatted_text(p, text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def main():
    files = [
        ("Combined_Manuscript_SciReports_2026.md", "Combined_Manuscript_SciReports_2026.docx"),
        ("Cover_Letter_SciReports_Revision_2026.md", "Cover_Letter_SciReports_Revision_2026.docx"),
        ("Response_to_Editor_SciReports_2026.md", "Response_to_Editor_SciReports_2026.docx"),
    ]

    results = []
    for md_name, docx_name in files:
        md_path = os.path.join(BASE, md_name)
        docx_path = os.path.join(BASE, docx_name)
        try:
            size = parse_markdown_to_word(md_path, docx_path)
            results.append(f"OK: {docx_name} ({size:.0f} KB)")
        except Exception as e:
            results.append(f"FAIL: {docx_name} - {e}")

    # Write status file
    status_path = os.path.join(BASE, "convert_status.txt")
    with open(status_path, "w") as f:
        for r in results:
            f.write(r + "\n")
    print("\n".join(results))


if __name__ == "__main__":
    main()
