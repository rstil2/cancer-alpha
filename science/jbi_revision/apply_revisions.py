#!/usr/bin/env python3
"""Apply proactive revisions to JBI manuscript (working copy only).

Reads from submitted_snapshot/, writes to working/.
Does NOT modify files in science/ root or submitted_snapshot/.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
SNAPSHOT = ROOT / "submitted_snapshot" / "Combined_Manuscript_JBI_SUBMITTED_2026-06-21.docx"
WORKING = ROOT / "working" / "Combined_Manuscript_JBI_REVISED.docx"

# Order matters for multi-word phrases (longer first).
REPLACEMENTS: list[tuple[str, str]] = [
    (
        "Experimental Design Dominates Model Architecture in Multi-Model Cancer Classification",
        "Experimental Design Dominates Model Architecture in Multi-Modal Cancer Classification",
    ),
    ("approaches perfect classification with 1,248 real balanced samples", "reaches 98.4% balanced accuracy on a held-out TCGA test set (n=250) with 1,248 real balanced training samples"),
    ("Study 2's near-perfect classification without synthetic augmentation", "Study 2's 98.4% balanced accuracy without synthetic augmentation"),
    ("4.7 Near-Perfect Classification from Large Balanced Real Data (Study 2)", "4.7 High-Accuracy Classification from Large Balanced Real Data (Study 2)"),
    ("Gradient boosting achieves clinical-grade accuracy with as few as 158 samples (SMOTE) and approaches perfect classification with 1,248 real balanced samples.", "Gradient boosting achieves high balanced accuracy with as few as 158 samples when SMOTE is applied (Study 1) and 98.4% balanced accuracy with 1,248 real balanced samples without synthetic augmentation (Study 2)."),
    ("gradient boosting consistently achieves clinical-grade multi-modal cancer classification", "gradient boosting consistently achieves high balanced accuracy for multi-modal cancer classification"),
    ("95.0 +/- 5.4% balanced accuracy-clinical-grade performance-while", "95.0 +/- 5.4% balanced accuracy—high performance in the small-n regime—while"),
    ("gradient boosting achieves clinical-grade accuracy with as few as 158 samples via SMOTE and approaches the Bayes error with 1,248 real balanced samples", "gradient boosting achieves 95.0% balanced accuracy with as few as 158 samples via SMOTE (Study 1) and 98.4% balanced accuracy with 1,248 real balanced samples (Study 2)"),
    ("clinical-grade performance from even smaller labelling budgets", "high balanced accuracy from even smaller labelling budgets"),
    ("Gradient boosting achieves clinical-grade accuracy with as few as 158 samples (SMOTE)", "Gradient boosting achieves high balanced accuracy with as few as 158 samples when SMOTE is applied (Study 1)"),
    ("KLK4", "KLK3"),
    (
        "In accordance with JAIR policy, we disclose that AI writing assistance",
        "In accordance with Elsevier and Journal of Biomedical Informatics policies, we disclose that AI writing assistance",
    ),
    (
        "The author has filed a provisional patent application (No. 63/847,316) related to methods described in this manuscript. No other competing interests exist.",
        "A provisional patent application (No. 63/847,316) was filed in 2024 but was not converted to a non-provisional application and has lapsed; no patent is currently in force. No other competing interests exist.",
    ),
    (
        "Source code: https://github.com/rstil2/cancer-alpha (src/pipeline/). TCGA raw data: GDC Data Portal (https://portal.gdc.cancer.gov/). ICGC ARGO data: ARGO platform (https://platform.icgc-argo.org/). Processed feature matrices, model predictions, and SHAP values will be deposited in Zenodo upon publication.",
        "Study 2 results are reproduced by the canonical pipeline at https://github.com/rstil2/cancer-alpha (src/pipeline/steps 1–4; see science/jbi_revision/supplementary/REPRODUCTION_GUIDE.md). Legacy demo and experimental scripts in the repository are not used for manuscript results. TCGA raw data: GDC Data Portal (https://portal.gdc.cancer.gov/). ICGC ARGO data: ARGO platform (https://platform.icgc-argo.org/). Processed feature matrices, model predictions, and SHAP values will be deposited in Zenodo upon acceptance.",
    ),
]

LIMITATIONS_OLD = (
    "Study 1's 158-sample dataset, while larger than typical rare disease cohorts, remains modest; "
    "preliminary external validation on 76 ICGC samples is encouraging but limited. Study 2 demonstrates "
    "deep learning underperforms at 1,248 samples, but extension to very large datasets (n>10,000) across "
    "33 cancer types may reveal regime transitions. More extensive validation across multiple independent "
    "institutions would substantially strengthen generalisability claims. Extension to survival prediction, "
    "treatment response, and rare cancers requires dedicated investigation. Integration of additional "
    "modalities (histopathology imaging, clinical notes) remains unexplored."
)

LIMITATIONS_NEW = (
    "Study 1's 158-sample dataset, while larger than typical rare disease cohorts, remains modest; "
    "external validation on 76 ICGC ARGO samples is encouraging but limited in size and feature overlap "
    "with the TCGA training set. Study 2 reports performance on a held-out TCGA test split (n=250); "
    "independent external cohorts (e.g., CPTAC, additional ICGC projects) were not evaluated and are needed "
    "before clinical deployment claims. Study 2 used stratified subsampling to 156 patients per cancer type "
    "(1,248 of 1,432 integrated samples); performance on the full imbalanced TCGA distribution may differ, "
    "and imbalance-robustness analyses are planned (see Section 3.8). The two studies used slightly different "
    "cancer-type panels (Study 1 included KIRC; Study 2 included LUSC instead) and feature sets (110 vs 4,063 "
    "features), so they should be interpreted as complementary regimes rather than a single longitudinal cohort. "
    "Study 1 employed SMOTE synthetic oversampling; Study 2 did not—synthetic augmentation remains appropriate "
    "only when real-data balancing is infeasible. Extension to cancer of unknown primary biopsy samples, survival "
    "prediction, treatment response, and rare cancers requires dedicated investigation. VC-dimension bounds are "
    "approximate and intended as interpretive support, not formal generalisation guarantees. Integration of "
    "additional modalities (histopathology imaging, clinical notes) remains unexplored."
)

IMBALANCE_METHODS = (
    "3.8 Sensitivity to Class Imbalance (Study 2 related). To assess whether balanced "
    "subsampling inflates performance, we trained LightGBM on a balanced cohort "
    "(150 samples per cancer type, 2,000 features) and evaluated on (1) a stratified "
    "balanced test set and (2) a test set resampled to natural TCGA prevalence "
    "(BRCA 30%, LUAD 18%, PRAD 15%, COAD 12%, LUSC 10%, HNSC 8%, STAD 4%, LIHC 3%). "
    "Balanced test balanced accuracy: 96.4%; imbalanced test: 95.7% (0.7 pp drop). "
    "See imbalance_stress_test.py and supplementary/imbalance_stress_test_results.json."
)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def replace_in_paragraph(paragraph: Paragraph, replacements: list[tuple[str, str]]) -> bool:
    text = paragraph.text
    if not text.strip():
        return False
    original = text
    for old, new in replacements:
        text = text.replace(old, new)
    if text != original:
        paragraph.clear()
        paragraph.add_run(text)
        return True
    return False


def insert_before(paragraph: Paragraph, text: str, bold_heading: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if bold_heading:
        new_para.add_run(text).bold = True
    else:
        new_para.add_run(text)
    return new_para


def apply_revisions() -> None:
    if not SNAPSHOT.exists():
        raise FileNotFoundError(f"Submitted snapshot not found: {SNAPSHOT}")

    WORKING.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SNAPSHOT, WORKING)

    doc = Document(str(WORKING))
    changed = 0

    for para in doc.paragraphs:
        if replace_in_paragraph(para, REPLACEMENTS):
            changed += 1
        if para.text.strip() == LIMITATIONS_OLD:
            para.clear()
            para.add_run(LIMITATIONS_NEW)
            changed += 1

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("4. Results") and i > 0:
            insert_paragraph_after(doc.paragraphs[i - 1], IMBALANCE_METHODS)
            changed += 1
            break

    for para in doc.paragraphs:
        if para.text.strip() == "Funding":
            insert_before(para, "R.C.S. conceived the study, designed the experiments, performed all analyses, interpreted results, and wrote the manuscript.")
            insert_before(para, "Author Contributions", bold_heading=True)
            insert_before(
                para,
                "This study used de-identified, publicly available data from TCGA and ICGC ARGO. TCGA data were collected under original project-specific institutional review board approvals with patient consent for research use. Secondary analysis of de-identified public genomic data does not require additional institutional review at the author's institution.",
            )
            insert_before(para, "Ethics Statement", bold_heading=True)
            changed += 1
            break

    doc.save(str(WORKING))
    print(f"Saved: {WORKING}")
    print(f"Edits applied: {changed}+")
    print("Review REVISION_CHANGELOG.md before any journal upload.")


if __name__ == "__main__":
    apply_revisions()
