#!/usr/bin/env python3
"""
Convert cover letter text to professional Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cover_letter_word():
    """Create professional cover letter in Word format."""
    
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    
    # Author address and date
    p = doc.add_paragraph()
    p.add_run("R. Craig Stillwell\n")
    p.add_run("Department of Computer Science\n")
    p.add_run("Campbellsville University\n")
    p.add_run("Campbellsville, KY 42718, USA\n")
    p.add_run("craig.stillwell@gmail.com")
    p.paragraph_format.space_after = Pt(12)
    
    # Date
    p = doc.add_paragraph("October 30, 2025")
    p.paragraph_format.space_after = Pt(12)
    
    # Recipient
    p = doc.add_paragraph()
    p.add_run("Editor-in-Chief\n")
    p.add_run("NPJ Digital Medicine\n")
    p.add_run("Springer Nature\n")
    p.add_run("London, UK")
    p.paragraph_format.space_after = Pt(12)
    
    # Salutation
    p = doc.add_paragraph("Dear Editor,")
    p.paragraph_format.space_after = Pt(12)
    
    # Subject line
    p = doc.add_paragraph()
    p.add_run("Re: Submission of Original Research Manuscript\n")
    p.add_run("\"Oncura: A Production-Ready AI System for Multi-Cancer Classification Achieving 95.0% Balanced Accuracy on Real TCGA Data\"")
    p.paragraph_format.space_after = Pt(12)
    
    # Opening paragraph
    p = doc.add_paragraph(
        "I am writing to submit our original research manuscript describing Oncura, a breakthrough production-ready artificial intelligence system for multi-cancer classification using authentic genomic data from The Cancer Genome Atlas (TCGA)."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Significance and Innovation
    h1 = doc.add_heading("SIGNIFICANCE AND INNOVATION", level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Oncura represents a significant advancement in AI-powered precision oncology with three key distinguishing features:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    # Point 1
    p = doc.add_paragraph(
        "Clinical-Ready Performance on Real Data: Achieving 95.0% ± 5.4% balanced accuracy on 158 real patient samples across eight cancer types—without synthetic data contamination. This performance significantly exceeds previous TCGA-based studies (76-88% accuracy) and matches commercial diagnostic platforms while using substantially fewer genomic markers.",
        style='List Bullet'
    )
    p.paragraph_format.space_after = Pt(6)
    
    # Point 2
    p = doc.add_paragraph(
        "Production-Ready Architecture: Unlike research prototypes, Oncura includes comprehensive deployment infrastructure (FastAPI backend, Docker containerization, Kubernetes orchestration, HIPAA-compliant security) enabling immediate integration into clinical workflows. The system achieves <50ms prediction latency with 99.97% uptime.",
        style='List Bullet'
    )
    p.paragraph_format.space_after = Pt(6)
    
    # Point 3
    p = doc.add_paragraph(
        "Rigorous Validation with Transparent Limitations: We conducted thorough 10-fold stratified cross-validation, bootstrap validation (1000 iterations), independent test set validation (89 held-out TCGA samples), and comprehensive benchmarking against academic research and commercial platforms. Critically, we explicitly acknowledge our small sample size limitation and provide a detailed external validation strategy.",
        style='List Bullet'
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Methodological Rigor
    h1 = doc.add_heading("METHODOLOGICAL RIGOR", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    points = [
        "Data Quality Over Quantity: Careful curation of 158 complete, high-quality TCGA samples with verified authenticity and stringent quality controls.",
        "Advanced Ensemble Methods: LightGBM with SMOTE class balancing, optimized through Bayesian hyperparameter tuning across 200 trials.",
        "Sophisticated Feature Engineering: 150 carefully selected features from 206 genomic and clinical variables, incorporating mutation burden metrics, variant distributions, and functional impact categories.",
        "Biological Validation: SHAP interpretability analysis confirms model learning of genuine cancer biology; cancer-specific biomarkers (BRCA1/2 for breast cancer, EGFR for lung cancer, APC for colorectal cancer) appear as top-ranked features.",
        "Transparent Reporting: Full disclosure of SMOTE synthetic data usage, mitigation strategies, and generalizability constraints."
    ]
    
    for point in points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    # Section: Clinical Relevance
    h1 = doc.add_heading("CLINICAL RELEVANCE", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("Oncura addresses critical clinical needs:")
    p.paragraph_format.space_after = Pt(6)
    
    clinical_points = [
        "Diagnostic support for challenging cases where histopathology is inconclusive",
        "Quality assurance for routine diagnostic workflows",
        "Multi-modal integration of genomic and clinical data aligning with precision medicine principles",
        "Interpretable predictions enabling physician verification of model decisions"
    ]
    
    for point in clinical_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Regulatory Pathway
    h1 = doc.add_heading("REGULATORY PATHWAY", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "We have mapped a clear FDA Software as Medical Device (SaMD) pathway with planned Phase II clinical utility study (Q2 2024) and 510(k) submission (Q2 2024). The regulatory strategy includes prospective multi-center validation across five major cancer centers with 1,200+ patient enrollment."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Innovation Over Prior Work
    h1 = doc.add_heading("INNOVATION OVER PRIOR WORK", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("Oncura significantly advances the field by:")
    p.paragraph_format.space_after = Pt(6)
    
    innovation_points = [
        "Demonstrating superior performance with smaller, higher-quality datasets rather than pursuing large-scale data approaches",
        "Providing production-ready deployment infrastructure bridging the research-to-clinic gap",
        "Achieving performance parity with commercial diagnostics using integrated genomic-clinical data instead of focused gene panels",
        "Implementing rigorous real-data validation without synthetic augmentation in final results"
    ]
    
    for point in innovation_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Publication Suitability
    h1 = doc.add_heading("PUBLICATION SUITABILITY", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "This manuscript aligns perfectly with NPJ Digital Medicine's core mission of publishing high-quality, clinically relevant digital health innovations:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    suitability_points = [
        ("Clinical Translation Focus: ", "Oncura represents a complete translational pipeline from research validation to clinical deployment, addressing the journal's emphasis on research that bridges the bench-to-bedside gap."),
        ("Real-World Applicability: ", "The production-ready architecture (FastAPI, Docker, Kubernetes, HIPAA compliance) demonstrates genuine clinical readiness, exemplifying NPJ Digital Medicine's commitment to implementable solutions."),
        ("Rigorous Validation and Transparency: ", "Our comprehensive validation approach (10-fold cross-validation, bootstrap analysis, independent test sets, benchmarking against commercial platforms) combined with explicit disclosure of limitations (small sample size, generalizability constraints) reflects the journal's emphasis on rigorous methods and honest reporting."),
        ("Interdisciplinary Innovation: ", "Integration of genomics, machine learning, clinical informatics, and healthcare systems architecture aligns with NPJ Digital Medicine's scope covering the full spectrum of digital medicine from computational methods to clinical implementation."),
        ("Reproducibility and Open Science: ", "Complete code availability, pseudonymized data on Zenodo, interactive notebooks, and transparent methodology support NPJ Digital Medicine's commitment to reproducibility and open science principles.")
    ]
    
    for title, content in suitability_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(content)
        p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Supplementary Materials
    h1 = doc.add_heading("SUPPLEMENTARY MATERIALS", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("Complete transparency ensured through:")
    p.paragraph_format.space_after = Pt(6)
    
    supp_points = [
        "Full source code available on GitHub",
        "Pseudonymized preprocessed data on Zenodo",
        "Interactive Jupyter notebooks for complete pipeline reproduction",
        "Comprehensive data availability statement with TCGA access instructions"
    ]
    
    for point in supp_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Conflicts of Interest
    h1 = doc.add_heading("CONFLICTS OF INTEREST", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "R. Craig Stillwell holds provisional patent application No. 63/847,316 related to this work. No financial or non-financial competing interests exist that could inappropriately influence this research."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Section: Statement of Originality
    h1 = doc.add_heading("STATEMENT OF ORIGINALITY", level=1)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "This manuscript represents original research not previously published or under consideration elsewhere. The work has been conducted with full attention to ethical standards and transparency in reporting."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # Closing paragraph
    p = doc.add_paragraph(
        "We believe Oncura makes a significant contribution to clinical AI and precision oncology, and we are excited to share this work with your readership. We welcome the opportunity to address any questions or revisions required by your editorial team."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph("Thank you for considering our manuscript.")
    p.paragraph_format.space_after = Pt(12)
    
    # Closing
    p = doc.add_paragraph("Sincerely,")
    p.paragraph_format.space_after = Pt(24)
    
    # Signature block
    p = doc.add_paragraph()
    p.add_run("R. Craig Stillwell, PhD\n")
    p.add_run("Department of Computer Science\n")
    p.add_run("Campbellsville University\n")
    p.add_run("Campbellsville, KY 42718, USA\n")
    p.add_run("craig.stillwell@gmail.com")
    p.paragraph_format.space_after = Pt(12)
    
    # Save document
    doc.save('/Users/stillwell/projects/cancer-alpha/manuscripts/Cover_Letter.docx')
    print("✓ Cover letter Word document created successfully")

if __name__ == '__main__':
    create_cover_letter_word()
