#!/usr/bin/env python3
"""
Convert markdown manuscript to Word document with embedded figures.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

def add_heading_style(doc, text, level):
    """Add a heading with appropriate style."""
    if level == 1:
        doc.add_heading(text, level=1)
    elif level == 2:
        doc.add_heading(text, level=2)
    elif level == 3:
        doc.add_heading(text, level=3)
    else:
        doc.add_heading(text, level=level)

def add_paragraph_with_citations(doc, text):
    """Add paragraph with proper formatting for citations."""
    # Replace markdown citation format [text](ref) with plain text
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\1', text)
    
    if text.strip():
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)

def embed_figure(doc, fig_num, fig_path, caption):
    """Embed a figure in the document."""
    if os.path.exists(fig_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path, width=Inches(6.5))
        
        # Add caption below figure
        caption_p = doc.add_paragraph(caption)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_format = caption_p.paragraph_format
        caption_format.line_spacing = 1.5
        caption_format.space_before = Pt(6)
        caption_format.space_after = Pt(12)
        
        # Make caption italicized
        for run in caption_p.runs:
            run.italic = True
            run.font.size = Pt(10)
        
        doc.add_paragraph()  # Add spacing after figure
    else:
        doc.add_paragraph(f"[Figure {fig_num}: {caption}]")

def convert_markdown_to_word(md_file, output_file):
    """Convert markdown manuscript to Word document."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    
    # Read markdown file
    with open(md_file, 'r') as f:
        content = f.read()
    
    # Split by main sections
    lines = content.split('\n')
    
    # Track figures to embed
    figure_map = {
        '1': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure1_model_performance.png', 
              'Figure 1: Model Performance Comparison. Comparison of balanced accuracy, precision, recall, and F1-score across six machine learning models.'),
        '2': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure2_cancer_type_performance.png',
              'Figure 2: Cancer Type-Specific Performance. Performance metrics across eight cancer types showing consistent accuracy above clinical threshold.'),
        '3': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure3_feature_importance.png',
              'Figure 3: Feature Importance Analysis. Top 20 features ranked by importance score with biological validation.'),
        '3b': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure3b_shap_interpretability.png',
               'Figure 3B: SHAP Summary Plot and Individual Force Plots. Global and patient-specific feature contributions to model predictions.'),
        '4': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure4_comparison_studies.png',
              'Figure 4: Benchmarking Analysis. Comparison with academic research and commercial platforms.'),
        '5a': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure5a_confusion_matrix.png',
               'Figure 5A: Confusion Matrix. Classification performance across all cancer types with minimal cross-type misclassification.'),
        '5b': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure5b_roc_curves.png',
               'Figure 5B: ROC Curves. Area-under-curve analysis showing consistently high performance across cancer types.'),
        '5': ('/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure6_system_architecture.png',
              'Figure 5: System Architecture. End-to-end workflow from data input through clinical deployment.'),
    }
    
    i = 0
    current_section = None
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines at beginning
        if not line and i < 5:
            i += 1
            continue
        
        # Title
        if line.startswith('# ') and not line.startswith('# '):
            title = line.replace('# ', '').strip()
            doc.add_heading(title, level=0)
            title_format = doc.paragraphs[-1].paragraph_format
            title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_format.space_after = Pt(12)
            i += 1
            continue
        
        # Author and affiliation
        if line.startswith('**R. Craig Stillwell'):
            doc.add_paragraph(line.replace('**', '').strip()).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            # Skip until we find Abstract
            while i < len(lines) and not lines[i].strip().startswith('## Abstract'):
                if lines[i].strip() and not lines[i].strip().startswith('*Corresponding') and not lines[i].strip().startswith('---'):
                    p = doc.add_paragraph(lines[i].strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            continue
        
        # Section headers (## )
        if line.startswith('## ') and not line.startswith('## '):
            section = line.replace('## ', '').strip()
            doc.add_heading(section, level=1)
            current_section = section
            i += 1
            continue
        
        # Subsection headers (### )
        if line.startswith('### ') and not line.startswith('### '):
            subsection = line.replace('### ', '').strip()
            doc.add_heading(subsection, level=2)
            i += 1
            continue
        
        # Sub-subsection headers (#### )
        if line.startswith('#### '):
            subsubsection = line.replace('#### ', '').strip()
            doc.add_heading(subsubsection, level=3)
            i += 1
            continue
        
        # Figure references - embed figures
        if line.startswith('**Figure '):
            fig_match = re.search(r'Figure (\d+[a-z]?)', line)
            if fig_match:
                fig_num = fig_match.group(1)
                if fig_num in figure_map:
                    fig_path, caption = figure_map[fig_num]
                    embed_figure(doc, fig_num, fig_path, caption)
            i += 1
            continue
        
        # Table handling
        if line.startswith('|'):
            # Simple table extraction
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse and create table
            rows = [l.split('|')[1:-1] for l in table_lines if l.strip()]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for row_idx, row in enumerate(rows):
                    for col_idx, cell in enumerate(row):
                        table.rows[row_idx].cells[col_idx].text = cell.strip()
                        if row_idx == 0:  # Header row
                            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            continue
        
        # Skip separator lines and empty lines
        if line.startswith('---') or not line:
            i += 1
            continue
        
        # Regular paragraph text
        if line:
            add_paragraph_with_citations(doc, line)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✓ Word document created: {output_file}")

if __name__ == '__main__':
    md_file = '/Users/stillwell/projects/cancer-alpha/manuscripts/cancer_alpha_main_manuscript.md'
    output_file = '/Users/stillwell/projects/cancer-alpha/manuscripts/Oncura_Manuscript.docx'
    
    convert_markdown_to_word(md_file, output_file)
