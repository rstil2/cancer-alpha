#!/usr/bin/env python3
"""
Convert corrected markdown manuscript to Word document with embedded figures.
Careful implementation with proper figure placement.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def add_paragraph_text(doc, text):
    """Add paragraph with proper formatting."""
    if text.strip():
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        return p

def embed_figure_in_doc(doc, fig_path, caption):
    """Embed figure at current position with caption."""
    if os.path.exists(fig_path):
        try:
            # Add centered figure
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(fig_path, width=Inches(6.0))
            
            # Add caption
            caption_p = doc.add_paragraph(caption)
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.line_spacing = 1.5
            caption_p.paragraph_format.space_before = Pt(6)
            caption_p.paragraph_format.space_after = Pt(12)
            for run in caption_p.runs:
                run.italic = True
                run.font.size = Pt(10)
            
            doc.add_paragraph()
            return True
        except Exception as e:
            print(f"Error embedding figure {fig_path}: {e}")
            return False
    else:
        print(f"Figure not found: {fig_path}")
        return False

def add_table_to_doc(doc, rows_data):
    """Add table with proper formatting."""
    if not rows_data:
        return
    
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Grid Accent 1'
    
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(cell_text).strip()
            
            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

def convert_to_word(md_file, output_file):
    """Convert markdown to Word document."""
    
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Figure mapping
    figures = {
        'figure1': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure1_model_performance.png',
        'figure2': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure2_cancer_type_performance.png',
        'figure3': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure3_feature_importance.png',
        'figure3b': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure3b_shap_interpretability.png',
        'figure4': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure4_comparison_studies.png',
        'figure5a': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure5a_confusion_matrix.png',
        'figure5b': '/Users/stillwell/projects/cancer-alpha/manuscripts/manuscript_figures/figure5b_roc_curves.png',
    }
    
    with open(md_file, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Title
        if line.startswith('# ') and i < 10:
            title = line.replace('# ', '').strip()
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Author information
        if line.startswith('**R. Craig Stillwell'):
            # Add author and affiliation
            while i < len(lines) and not lines[i].strip().startswith('## Abstract'):
                if lines[i].strip() and not lines[i].strip().startswith('---'):
                    p = doc.add_paragraph(lines[i].strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
            continue
        
        # Section headers
        if line.startswith('## ') and not line.startswith('## '):
            section = line.replace('## ', '').strip()
            doc.add_heading(section, level=1)
            i += 1
            continue
        
        # Subsection headers
        if line.startswith('### ') and not line.startswith('### '):
            subsection = line.replace('### ', '').strip()
            doc.add_heading(subsection, level=2)
            i += 1
            continue
        
        # Sub-subsection headers
        if line.startswith('#### '):
            subsubsection = line.replace('#### ', '').strip()
            doc.add_heading(subsubsection, level=3)
            i += 1
            continue
        
        # Handle Figure 1
        if 'Figure 1: Model Performance Comparison' in line:
            add_paragraph_text(doc, line.replace('**Figure 1: Model Performance Comparison**', '').strip())
            i += 1
            if i < len(lines) and lines[i].strip().startswith('Comparison of'):
                caption = "Figure 1: Model Performance Comparison. " + lines[i].strip()
                embed_figure_in_doc(doc, figures['figure1'], caption)
                i += 1
            continue
        
        # Handle Figure 2
        if 'Figure 2: Cancer Type-Specific Performance' in line:
            i += 1
            if i < len(lines) and lines[i].strip().startswith('Performance metrics'):
                caption = "Figure 2: Cancer Type-Specific Performance. " + lines[i].strip()
                embed_figure_in_doc(doc, figures['figure2'], caption)
                i += 1
            continue
        
        # Handle Figure 3
        if 'Figure 3: Feature Importance Analysis' in line:
            i += 1
            caption = "Figure 3: Feature Importance Analysis. Top 20 features ranked by importance score with biological validation."
            embed_figure_in_doc(doc, figures['figure3'], caption)
            continue
        
        # Handle Figure 3B
        if 'Figure 3B: SHAP' in line:
            i += 1
            caption = "Figure 3B: SHAP Summary Plot and Individual Force Plots. Global and patient-specific feature contributions to model predictions."
            embed_figure_in_doc(doc, figures['figure3b'], caption)
            continue
        
        # Handle Figure 4
        if 'Figure 4: Benchmarking Analysis' in line:
            i += 1
            if i < len(lines) and lines[i].strip().startswith('Comparison with'):
                caption = "Figure 4: Benchmarking Analysis. " + lines[i].strip()
                embed_figure_in_doc(doc, figures['figure4'], caption)
                i += 1
            continue
        
        # Handle confusion matrix and ROC figures
        if 'Figure 5A' in line or 'Figure 5B' in line:
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            rows = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                if cells and '---' not in cells[0]:
                    rows.append(cells)
            
            if rows:
                add_table_to_doc(doc, rows)
                doc.add_paragraph()
            continue
        
        # Skip separator lines
        if line.startswith('---') or not line:
            i += 1
            continue
        
        # Regular text
        if line:
            add_paragraph_text(doc, line)
        
        i += 1
    
    doc.save(output_file)
    print(f"✓ Word document created: {output_file}")

if __name__ == '__main__':
    md_file = '/Users/stillwell/projects/cancer-alpha/manuscripts/cancer_alpha_main_manuscript.md'
    output_file = '/Users/stillwell/projects/cancer-alpha/manuscripts/Oncura_Manuscript_Final.docx'
    
    convert_to_word(md_file, output_file)
