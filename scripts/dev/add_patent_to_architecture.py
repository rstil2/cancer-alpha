#!/usr/bin/env python3
"""
Add patent information to the architecture paper Word document and convert to PDF.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import subprocess
import os

def add_patent_info_to_docx():
    """Add patent information to the existing Word document."""
    
    # Open the existing Word document
    doc = Document('preprints/multimodal_transformer_architecture_corrected.docx')
    
    # Find the title paragraph (first paragraph)
    title_paragraph = doc.paragraphs[0]
    
    # Insert patent warning after the title
    patent_para = doc.paragraphs[1] if len(doc.paragraphs) > 1 else doc.add_paragraph()
    
    # Clear existing content if any
    patent_para.clear()
    
    # Add patent warning
    patent_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the patent information
    run1 = patent_para.add_run("⚠️ PATENT PROTECTED TECHNOLOGY ⚠️\n\n")
    run1.bold = True
    run1.font.size = Inches(0.14)  # 14pt
    
    run2 = patent_para.add_run("Patent: ")
    run2.bold = True
    run3 = patent_para.add_run("Provisional Application No. 63/847,316\n")
    
    run4 = patent_para.add_run("Title: ")
    run4.bold = True
    run5 = patent_para.add_run("Systems and Methods for Cancer Classification Using Multi-Modal Transformer-Based Architectures\n")
    
    run6 = patent_para.add_run("Patent Holder: ")
    run6.bold = True
    run7 = patent_para.add_run("Dr. R. Craig Stillwell\n")
    
    run8 = patent_para.add_run("Commercial Use: ")
    run8.bold = True
    run9 = patent_para.add_run("Requires separate patent license\n\n")
    
    run10 = patent_para.add_run("For commercial licensing inquiries, contact: craig.stillwell@gmail.com")
    run10.italic = True
    
    # Add horizontal line after patent info
    patent_para.add_run("\n" + "—" * 80 + "\n")
    
    # Save the modified document
    doc.save('preprints/multimodal_transformer_architecture_corrected.docx')
    print("✅ Added patent information to Word document")
    
    return True

def convert_docx_to_pdf():
    """Convert the Word document to PDF."""
    
    try:
        # Try using LibreOffice to convert
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', 
            '--outdir', 'preprints/',
            'preprints/multimodal_transformer_architecture_corrected.docx'
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            print("✅ Successfully converted to PDF using LibreOffice")
            return True
        else:
            print(f"LibreOffice conversion failed: {result.stderr}")
            
    except FileNotFoundError:
        print("LibreOffice not found, trying alternative...")
    
    try:
        # Try using pandoc
        result = subprocess.run([
            'pandoc', 
            'preprints/multimodal_transformer_architecture_corrected.docx',
            '-o', 'preprints/multimodal_transformer_architecture_corrected.pdf'
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            print("✅ Successfully converted to PDF using pandoc")
            return True
        else:
            print(f"Pandoc conversion failed: {result.stderr}")
            
    except FileNotFoundError:
        print("Pandoc not found")
    
    print("❌ Could not convert to PDF automatically")
    print("Please manually open the Word document and save as PDF")
    return False

if __name__ == "__main__":
    try:
        # Add patent information
        add_patent_info_to_docx()
        
        # Convert to PDF
        convert_docx_to_pdf()
        
    except ImportError:
        print("❌ python-docx not found. Installing...")
        subprocess.run(["pip", "install", "python-docx"])
        add_patent_info_to_docx()
        convert_docx_to_pdf()
        
    except Exception as e:
        print(f"❌ Error: {e}")
        print("The Word document has been updated with patent information.")
        print("Please manually convert it to PDF if automatic conversion failed.")
